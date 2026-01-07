"""
DOCX extractor using python-docx.

Extracts Word document content including paragraphs, headings, and tables.
"""

import io
from typing import Any

from docx import Document

from .base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    """Extract content from DOCX files.

    Features:
    - Extract document properties (title, author, dates)
    - Extract paragraphs with heading hierarchy
    - Convert tables to markdown tables
    - Preserve bold, italic formatting
    """

    def extract(self, content: bytes, filename: str) -> ExtractionResult:
        """Extract text content from DOCX file bytes.

        Args:
            content: Raw DOCX content as bytes.
            filename: Original filename.

        Returns:
            ExtractionResult with markdown content and metadata.
        """
        title = self._extract_title_from_filename(filename)

        # Read DOCX
        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            # Fall back to plain text with warning
            return self._create_fallback_result(content, filename, title, str(e))

        # Extract document properties
        doc_title = doc.core_properties.title or title
        author = doc.core_properties.author
        created = doc.core_properties.created
        modified = doc.core_properties.modified

        # Count elements
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)

        # Extract content
        markdown_body = self._extract_content(doc)
        word_count = self._count_words(markdown_body)

        # Build structural metadata
        structural_metadata: dict[str, Any] = {
            "title": doc_title,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
        }
        if author:
            structural_metadata["author"] = author
        if created:
            structural_metadata["created"] = created.isoformat()
        if modified:
            structural_metadata["modified"] = modified.isoformat()

        # Build frontmatter metadata
        frontmatter_metadata = {
            "source_file": filename,
            "file_type": "docx",
            "title": doc_title,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
        }
        if author:
            frontmatter_metadata["author"] = author

        frontmatter = self._generate_frontmatter(frontmatter_metadata)

        # Combine frontmatter and content
        markdown = f"{frontmatter}\n{markdown_body}"

        return ExtractionResult(
            markdown=markdown,
            file_type="docx",
            title=doc_title,
            word_count=word_count,
            structural_metadata=structural_metadata,
            parse_warning=None,
        )

    def _extract_content(self, doc: Document) -> str:
        """Extract content from DOCX as markdown."""
        lines = []

        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element == element:
                        md = self._paragraph_to_markdown(para)
                        if md:
                            lines.append(md)
                        break

            # Check if it's a table
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == element:
                        md = self._table_to_markdown(table)
                        if md:
                            lines.append(md)
                        break

        return "\n\n".join(lines)

    def _paragraph_to_markdown(self, para) -> str:
        """Convert a paragraph to markdown."""
        text = para.text.strip()
        if not text:
            return ""

        # Check for heading style
        style_name = para.style.name if para.style else ""

        if style_name.startswith("Heading"):
            # Extract heading level
            try:
                level = int(style_name.replace("Heading", "").strip())
            except ValueError:
                level = 1
            return "#" * level + " " + text

        if style_name == "Title":
            return "# " + text

        # Regular paragraph - extract formatting
        formatted_text = self._format_runs(para.runs)
        return formatted_text if formatted_text else text

    def _format_runs(self, runs) -> str:
        """Format paragraph runs with bold/italic."""
        parts = []
        for run in runs:
            text = run.text
            if not text:
                continue

            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            parts.append(text)

        return "".join(parts)

    def _table_to_markdown(self, table) -> str:
        """Convert a table to markdown format."""
        lines = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

            # Add header separator after first row
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        return "\n".join(lines)

    def _create_fallback_result(
        self, content: bytes, filename: str, title: str, error: str
    ) -> ExtractionResult:
        """Create fallback result when DOCX parsing fails."""
        text = self._decode_content(content)
        word_count = self._count_words(text)
        frontmatter_metadata = {
            "source_file": filename,
            "file_type": "docx",
            "parse_warning": f"DOCX parsing failed: {error}",
        }
        frontmatter = self._generate_frontmatter(frontmatter_metadata)
        markdown = f"{frontmatter}\nFailed to parse DOCX: {error}"

        return ExtractionResult(
            markdown=markdown,
            file_type="docx",
            title=title,
            word_count=word_count,
            structural_metadata={},
            parse_warning=f"DOCX parsing failed: {error}",
        )
