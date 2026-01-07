"""Unit tests for DOCX extractor."""

import io

import pytest

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    Document = None

from ragstack_common.text_extractors.docx_extractor import DocxExtractor
from ragstack_common.text_extractors.base import ExtractionResult


def create_minimal_docx(
    title: str | None = None,
    paragraphs: list[str] | None = None,
    headings: list[tuple[str, int]] | None = None,
    table_data: list[list[str]] | None = None,
) -> bytes:
    """Create a minimal DOCX file for testing."""
    if Document is None:
        pytest.skip("python-docx not installed")

    doc = Document()

    # Set title in core properties if provided
    if title:
        doc.core_properties.title = title

    # Add headings
    if headings:
        for text, level in headings:
            doc.add_heading(text, level=level)

    # Add paragraphs
    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    # Add table
    if table_data:
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text

    # Write to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest.fixture
def simple_docx() -> bytes:
    """Fixture for simple DOCX with paragraphs."""
    return create_minimal_docx(
        title="Test Document",
        paragraphs=["First paragraph of content.", "Second paragraph of content."],
    )


@pytest.fixture
def docx_with_headings() -> bytes:
    """Fixture for DOCX with headings."""
    return create_minimal_docx(
        title="Headings Document",
        headings=[
            ("Main Title", 0),
            ("Section One", 1),
            ("Subsection", 2),
        ],
        paragraphs=["Some content under the headings."],
    )


@pytest.fixture
def docx_with_table() -> bytes:
    """Fixture for DOCX with table."""
    return create_minimal_docx(
        title="Table Document",
        paragraphs=["Introduction text."],
        table_data=[
            ["Header 1", "Header 2", "Header 3"],
            ["Cell 1", "Cell 2", "Cell 3"],
            ["Cell 4", "Cell 5", "Cell 6"],
        ],
    )


@pytest.fixture
def corrupted_docx() -> bytes:
    """Fixture for corrupted DOCX (invalid content)."""
    return b"not a valid docx file"


class TestDocxExtractor:
    """Tests for DocxExtractor."""

    def test_extracts_simple_docx(self, simple_docx):
        """Test extraction of simple DOCX."""
        extractor = DocxExtractor()
        result = extractor.extract(simple_docx, "document.docx")

        assert isinstance(result, ExtractionResult)
        assert result.file_type == "docx"

    def test_extracts_title_from_properties(self, simple_docx):
        """Test extraction of title from document properties."""
        extractor = DocxExtractor()
        result = extractor.extract(simple_docx, "document.docx")

        assert result.title == "Test Document"

    def test_extracts_paragraphs(self, simple_docx):
        """Test extraction of paragraph content."""
        extractor = DocxExtractor()
        result = extractor.extract(simple_docx, "document.docx")

        assert "First paragraph" in result.markdown
        assert "Second paragraph" in result.markdown

    def test_extracts_headings(self, docx_with_headings):
        """Test extraction of headings."""
        extractor = DocxExtractor()
        result = extractor.extract(docx_with_headings, "headings.docx")

        assert "Main Title" in result.markdown
        assert "Section One" in result.markdown

    def test_converts_headings_to_markdown(self, docx_with_headings):
        """Test that headings are converted to markdown format."""
        extractor = DocxExtractor()
        result = extractor.extract(docx_with_headings, "headings.docx")

        # Should have markdown heading markers
        assert "#" in result.markdown

    def test_extracts_tables(self, docx_with_table):
        """Test extraction of table content."""
        extractor = DocxExtractor()
        result = extractor.extract(docx_with_table, "table.docx")

        assert "Header 1" in result.markdown
        assert "Cell 1" in result.markdown
        # Should have markdown table format
        assert "|" in result.markdown

    def test_generates_frontmatter(self, simple_docx):
        """Test that frontmatter is generated correctly."""
        extractor = DocxExtractor()
        result = extractor.extract(simple_docx, "document.docx")

        assert result.markdown.startswith("---\n")
        assert "source_file: document.docx" in result.markdown
        assert "file_type: docx" in result.markdown

    def test_structural_metadata_includes_paragraph_count(self, simple_docx):
        """Test structural metadata includes paragraph count."""
        extractor = DocxExtractor()
        result = extractor.extract(simple_docx, "document.docx")

        assert "paragraph_count" in result.structural_metadata
        assert result.structural_metadata["paragraph_count"] >= 2

    def test_structural_metadata_includes_table_count(self, docx_with_table):
        """Test structural metadata includes table count."""
        extractor = DocxExtractor()
        result = extractor.extract(docx_with_table, "table.docx")

        assert "table_count" in result.structural_metadata
        assert result.structural_metadata["table_count"] == 1

    def test_corrupted_docx_falls_back(self, corrupted_docx):
        """Test that corrupted DOCX falls back with warning."""
        extractor = DocxExtractor()
        result = extractor.extract(corrupted_docx, "broken.docx")

        assert isinstance(result, ExtractionResult)
        assert result.parse_warning is not None

    def test_title_fallback_to_filename(self):
        """Test title falls back to filename when no title in properties."""
        docx_bytes = create_minimal_docx(paragraphs=["Just content."])
        extractor = DocxExtractor()
        result = extractor.extract(docx_bytes, "my_document.docx")

        # Should use filename when no title property
        assert result.title is not None


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
