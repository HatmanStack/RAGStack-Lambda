#!/usr/bin/env python3

"""
Generate sample test documents for RAGStack-Lambda testing.

Requires:
- PyMuPDF (pip install pymupdf)
- Pillow (pip install Pillow)
- python-docx (pip install python-docx)
- openpyxl (pip install openpyxl)
"""

from pathlib import Path


def generate_text_native_pdf():
    """Generate text-native PDF with embedded text."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        print("⚠️  PyMuPDF not installed. Install with: pip install pymupdf")
        return False

    print("Creating text-native.pdf...")

    doc = fitz.open()
    page = doc.new_page()

    # Add text content
    text_content = """Sample Text-Native PDF

This is a test document with embedded text.
It should be processed using direct text extraction, not OCR.

Date: 2025-01-15
Status: Test Document
Type: Text-Native PDF

Content:
This document contains regular text that is embedded in the PDF format.
When processed by RAGStack-Lambda, it should be detected as text-native
and skip the OCR step entirely.

Features tested:
- Direct text extraction
- Fast processing time (< 5 seconds)
- Cost optimization (no OCR charges)
- Page rendering for image embeddings

Additional Information:
This multi-page document tests pagination and chunking.
Each page should be processed sequentially, and text should be
extracted efficiently without requiring OCR services.

Expected Results:
- is_text_native: true
- total_pages: 1
- processing_time: < 5 seconds
- ocr_cost: $0.00
- embedding_cost: ~$0.001

Test Scenario:
Upload this file via the UI and verify it processes correctly.
Check the dashboard to confirm text-native detection worked.
"""

    # Insert text at position (50, 50)
    page.insert_text((50, 50), text_content, fontsize=11)

    # Save PDF
    doc.save("text-native.pdf")
    doc.close()

    print("✓ Created text-native.pdf")
    return True


def generate_invoice_image():
    """Generate sample invoice image."""
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print("⚠️  Pillow not installed. Install with: pip install Pillow")
        return False

    print("Creating invoice.jpg...")

    # Create image
    img = Image.new("RGB", (800, 1000), color="white")
    draw = ImageDraw.Draw(img)

    # Try to use default font
    try:
        font_large = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 24)
        font_normal = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 14)
    except OSError:
        font_large = ImageFont.load_default()
        font_normal = ImageFont.load_default()

    # Draw invoice content
    y = 50

    # Header
    draw.text((50, y), "INVOICE", fill="black", font=font_large)
    y += 60

    # Invoice details
    invoice_text = [
        "Invoice #: INV-2025-001",
        "Date: January 15, 2025",
        "Due Date: February 15, 2025",
        "",
        "Bill To:",
        "Acme Corporation",
        "123 Main Street",
        "Springfield, USA 12345",
        "",
        "Items:",
        "",
        "1. Professional Services         $1,500.00",
        "2. Cloud Infrastructure          $750.00",
        "3. Technical Support             $250.00",
        "",
        "                   Subtotal:     $2,500.00",
        "                   Tax (10%):    $250.00",
        "                   Total:        $2,750.00",
        "",
        "Payment Terms: Net 30 days",
        "Thank you for your business!",
    ]

    for line in invoice_text:
        draw.text((50, y), line, fill="black", font=font_normal)
        y += 25

    # Save image
    img.save("invoice.jpg", "JPEG", quality=95)

    print("✓ Created invoice.jpg")
    return True


def generate_word_document():
    """Generate sample Word document."""
    try:
        from docx import Document
    except ImportError:
        print("⚠️  python-docx not installed. Install with: pip install python-docx")
        return False

    print("Creating document.docx...")

    doc = Document()

    # Add title
    doc.add_heading("Sample Word Document", 0)

    # Add content
    doc.add_paragraph("This is a sample Microsoft Word document for testing RAGStack-Lambda.")

    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document tests the format conversion pipeline. "
        "Word documents should be converted to PDF format before OCR processing."
    )

    doc.add_heading("Test Criteria", level=1)
    doc.add_paragraph("1. Document should be converted to PDF")
    doc.add_paragraph("2. Text should be extracted correctly")
    doc.add_paragraph("3. Formatting should be preserved where possible")

    doc.add_heading("Expected Results", level=1)
    doc.add_paragraph("• File format: DOCX → PDF")
    doc.add_paragraph("• Text extraction: Successful")
    doc.add_paragraph("• Processing time: < 15 seconds")

    # Save document
    doc.save("document.docx")

    print("✓ Created document.docx")
    return True


def generate_excel_spreadsheet():
    """Generate sample Excel spreadsheet."""
    try:
        from openpyxl import Workbook
    except ImportError:
        print("⚠️  openpyxl not installed. Install with: pip install openpyxl")
        return False

    print("Creating spreadsheet.xlsx...")

    wb = Workbook()
    ws = wb.active
    ws.title = "Sales Data"

    # Add headers
    headers = ["Date", "Product", "Quantity", "Price", "Total"]
    ws.append(headers)

    # Add data
    data = [
        ["2025-01-01", "Widget A", 10, 25.00, 250.00],
        ["2025-01-02", "Widget B", 5, 50.00, 250.00],
        ["2025-01-03", "Widget C", 15, 10.00, 150.00],
        ["2025-01-04", "Widget A", 20, 25.00, 500.00],
        ["2025-01-05", "Widget B", 8, 50.00, 400.00],
    ]

    for row in data:
        ws.append(row)

    # Add totals
    ws.append([])
    ws.append(["", "", "Total", "", "=SUM(E2:E6)"])

    # Save spreadsheet
    wb.save("spreadsheet.xlsx")

    print("✓ Created spreadsheet.xlsx")
    return True


def main():
    print("=" * 60)
    print("RAGStack-Lambda Sample Document Generator")
    print("=" * 60)
    print()

    # Change to script directory
    script_dir = Path(__file__).parent
    script_dir.mkdir(parents=True, exist_ok=True)
    import os

    os.chdir(script_dir)

    results = {
        "text-native.pdf": generate_text_native_pdf(),
        "invoice.jpg": generate_invoice_image(),
        "document.docx": generate_word_document(),
        "spreadsheet.xlsx": generate_excel_spreadsheet(),
    }

    print()
    print("=" * 60)
    print("Summary:")
    print("=" * 60)

    success_count = sum(results.values())
    total_count = len(results)

    for filename, success in results.items():
        status = "✓" if success else "✗"
        print(f"{status} {filename}")

    print()
    print(f"Created {success_count}/{total_count} files successfully")

    if success_count < total_count:
        print()
        print("Install missing dependencies with:")
        print("  pip install pymupdf Pillow python-docx openpyxl")

    print()
    print("Note: For scanned.pdf, you'll need to provide an actual scanned document")
    print("or use an online tool to convert an image to a scanned PDF.")


if __name__ == "__main__":
    main()
